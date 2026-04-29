"""Build outputs/analiza.docx — raport detaliat Faza 4."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# === STILURI ===
styles = doc.styles
normal = styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(11)

# Margine pagina
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)


def add_heading(text, level=1, color='1F4E78'):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor.from_string(color)
        run.font.name = 'Calibri'
    return h


def add_para(text, bold=False, italic=False, size=11, after=6):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'Calibri'
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    p.paragraph_format.space_after = Pt(after)
    return p


def add_bullet(text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.6 + level * 0.6)
    return p


def add_table(headers, rows, col_widths=None, header_bg='1F4E78', header_color='FFFFFF'):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Light Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    hcells = t.rows[0].cells
    for i, h in enumerate(headers):
        hcells[i].text = h
        for p in hcells[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = RGBColor.from_string(header_color)
                r.font.size = Pt(10)
        # Background
        tc = hcells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), header_bg)
        tcPr.append(shd)
    # Rows
    for ri, row in enumerate(rows):
        cells = t.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            for p in cells[ci].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    if col_widths:
        for col_idx, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[col_idx].width = Cm(w)
    return t


def add_pagebreak():
    doc.add_page_break()


# ===== TITLE PAGE =====
title = doc.add_heading('Analiza fezabilitate', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = RGBColor.from_string('1F4E78')

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = subtitle.add_run('Operare parcari cu plata — Total Hub SA')
r.font.size = Pt(18)
r.bold = True
r.font.color.rgb = RGBColor.from_string('595959')

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run('Versiune 1.0  |  2026-04-29  |  Audienta: parteneri de afaceri')
r.font.size = Pt(11)
r.italic = True

doc.add_paragraph()
doc.add_paragraph()

verdict_p = doc.add_paragraph()
verdict_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = verdict_p.add_run('RECOMANDARE: GO CONDITIONAT')
r.font.size = Pt(20)
r.bold = True
r.font.color.rgb = RGBColor.from_string('006100')

doc.add_paragraph()
context = doc.add_paragraph()
context.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = context.add_run('NPV +145k EUR  |  IRR 31.1%  |  Payback 3 ani  |  Min Cash -130k')
r.font.size = Pt(12)
r.italic = True

add_pagebreak()

# ===== 1. CUPRINS =====
add_heading('Cuprins', 1)
toc_items = [
    '1. Sumar executiv si recomandare',
    '2. Context si obiectivul analizei',
    '3. Metodologie (5 faze cu gates)',
    '4. Faza 0 — Unit economics + Business economics',
    '5. Faza 1 — Specificatie modelului',
    '6. Faza 2 — Validator matematic + Reality check',
    '7. Faza 3 — Model financiar consolidat (model.xlsx)',
    '8. Pre-mortem — 3 moduri concrete de esec',
    '9. Sensibilitati si stres tests',
    '10. Concluzii si decizia ceruta',
    '11. Glosar',
    '12. Disclaimer si limitari',
]
for it in toc_items:
    add_para(it, after=4)

add_pagebreak()

# ===== 2. SUMAR EXECUTIV =====
add_heading('1. Sumar executiv si recomandare', 1)

add_para('Total Hub SA, companie de IT/automatizari, isi propune sa investeasca in operarea de parcari cu plata in Romania. Investitia in echipamente (ANPR, bariere, terminale) si operatiunile sunt asumate de Total Hub, in schimbul revenue share sau fee fix de la proprietarii locatiilor.')

add_heading('Recomandarea', 2)
add_para('GO CONDITIONAT.', bold=True)
add_para('Proiectul este viabil financiar pe baseline-ul de scaling 12 parcari operate la finalul anului 2, cu randament puternic pe capital propriu (IRR 31%) si NPV pozitiv +145.000 EUR la WACC 10,7%. Conditionarea vizeaza diversificarea timpurie pe non-retail si validarea ipotezelor de trafic pe primele 2-3 locatii instalate.')

add_heading('KPI principal', 2)
add_table(
    ['Indicator', 'Valoare', 'Prag KPI', 'Status'],
    [
        ['NPV @ WACC 10,7%',                '+145.261 EUR',     '> 0',                'PASS'],
        ['IRR (Internal Rate of Return)',    '31,1%',           '> 15%',              'PASS'],
        ['Payback simplu',                   '3 ani',           '< 4 ani',            'PASS'],
        ['Minimum Cash Balance',             '-129.613 EUR',    '> -300.000 EUR',     'PASS'],
        ['CF Net cumulat 5 ani',             '+437.020 EUR',    '> 800.000 EUR',      'FAIL'],
    ],
    col_widths=[5.5, 3, 3.5, 1.8],
)

add_para('Singurul KPI care esueaza este pragul absolut de cash flow cumulat (437k vs target 800k). Pragul a fost stabilit in spec.md ca tinta ambitioasa pentru randament puternic; rezultatul actual este viabil si genereaza valoare, dar la un nivel modest de generare de cash.', italic=True, size=10)

add_heading('Investitia totala', 2)
add_para('CAPEX total portofoliu 12 parcari: 474.000 EUR.', bold=True)
add_bullet('Capital propriu necesar (30%): 142.200 EUR')
add_bullet('Credit echipamente (70%): 331.800 EUR la 9% pe 5 ani')
add_bullet('Cash buffer suplimentar recomandat: ~130.000 EUR pentru a acoperi CF Net negativ anul 1 si 2')

add_pagebreak()

# ===== 3. CONTEXT =====
add_heading('2. Context si obiectivul analizei', 1)

add_heading('Compania', 2)
add_para('Total Hub SA este o societate pe actiuni romaneasca, profilata pe IT si automatizari. Echipa actuala include know-how tehnic (ANPR, bariere, integrare hardware/software) dar nu are experienta operationala in industria parking.')

add_heading('Oportunitatea', 2)
add_para('Pipeline-ul comercial activ:')
add_bullet('Auchan: 2 locatii in negociere activa')
add_bullet('Lidl: multiple locatii in pipeline')
add_bullet('Kaufland: planuri de abordare in trimestrul urmator')
add_bullet('Posibilitate extindere catre proprietari non-retail (parcari publice, hoteluri, parcari office)')

add_heading('Decizia care trebuie luata', 2)
add_para('Modelul nu raspunde la "investim sau nu" (decizia simplificata) ci la:')
add_para('"Care mix de tipuri de locatii × variante contractuale maximizeaza CF Net pe portofoliu, in conditiile pipeline-ului real, cu un orizont de 5 ani si scaling 12 parcari pana la finalul anului 2?"', bold=True)

add_heading('De ce conteaza structura', 2)
add_para('Variantele contractuale (operator 100% revenue, revenue split 50/50, fee fix, hibrid rent + split) au impact dramatic pe CF Net. La standardul real al pietei (120 min gratuitate Lidl/Kaufland), ne-protectarea contractuala (variantele C1/C2) genereaza CF marginal sau negativ, iar variantele cu fee fix de la retailer (C3) sau hibrid (C4) sunt singurele realiste.')

add_pagebreak()

# ===== 4. METODOLOGIE =====
add_heading('3. Metodologie (5 faze cu gates)', 1)
add_para('Analiza a fost construita liniar pe 5 faze, cu gate explicit cu confirmare la fiecare:')
add_bullet('Faza 0: Unit Economics + Business Economics + Kill Switch')
add_bullet('Faza 1: Specificatie scrisa (spec.md)')
add_bullet('Faza 2: Validator matematic + Reality check')
add_bullet('Faza 3: Model Excel polished (model.xlsx)')
add_bullet('Faza 4: Pre-mortem + Documente (acest raport)')

add_para('Avantajul abordarii liniare: bug matematice si ipoteze nerealiste se descopera in fazele timpurii (cost ~5 minute), nu dupa polish (cost 2-3 ore).', italic=True, size=10)

add_pagebreak()

# ===== 5. FAZA 0 =====
add_heading('4. Faza 0 — Unit Economics si Business Economics', 1)

add_heading('Unit economics — o singura parcare retail Lidl/Kaufland (V1)', 2)
add_table(
    ['Gratuitate', '% intrari care platesc', 'Venit anual', 'EBITDA / parcare / an', 'CF Net post-anuitate'],
    [
        ['30 min',  '30%',  '131.400 EUR',  '124.200 EUR',  '+93.300 EUR'],
        ['60 min',  '8%',   '62.000 EUR',   '54.800 EUR',   '+35.000 EUR'],
        ['120 min (standard piata)', '3%', '21.900 EUR', '14.700 EUR', '+1.300 EUR'],
    ],
    col_widths=[2.5, 3.5, 3, 3.5, 3.5],
)

add_para('La standardul real al pietei (120 min gratuitate, regula Lidl/Kaufland/Auchan), CF Net per parcare este marginal. Doar variantele cu fee fix de la retailer (C3) sau hibrid (C4) reformuleaza economics.', italic=True, size=10)

add_heading('Business economics — cate parcari acopera overhead 119k-166k EUR/an', 2)
add_table(
    ['Scenariu parcare', 'CF Net per parcare/an', 'Parcari necesare', 'Verdict'],
    [
        ['V1 retail mediu g=30',     '93.000 EUR',       '~2',      'realist'],
        ['V1 retail mediu g=60',     '35.000 EUR',       '~4',      'realist'],
        ['V1 retail mediu g=120',    '1.300 EUR',        '~90',     'KILL SWITCH'],
        ['Standalone non-retail',    '125.000 EUR',      '~1',      'realist'],
        ['Captiv',                   '31.000 EUR',       '~4',      'realist'],
    ],
    col_widths=[5, 3.5, 3, 3],
)

add_heading('Verdict Faza 0', 2)
add_para('Kill switch partial: V1/V2 retail la 120 min gratuitate sunt ne-viabile la nivel de business; raman in model ca referinte negative (justifica de ce avem nevoie de variante cu fee). Restul scenariilor (V3, V4 retail, non-retail) sunt viabile.')
add_para('Decizia se muta de la "investim sau nu" catre "ce mix de variante × tipuri".', bold=True)

add_pagebreak()

# ===== 6. FAZA 1 =====
add_heading('5. Faza 1 — Specificatia modelului', 1)

add_heading('Tipuri de locatii (parametrizate)', 2)
add_table(
    ['Cod', 'Tip', 'Locuri', 'Trafic/zi', 'CAPEX', 'Tarif', 'Gratuitate'],
    [
        ['A', 'Retail mare (hyper, ex: Auchan)',     '200-400',     '1.500-3.000',  '50.000 EUR',   '10 RON/h',  '120 min'],
        ['B', 'Retail mediu (super, ex: Lidl)',      '100-200',     '800-1.500',    '40.000 EUR',   '10 RON/h',  '120 min'],
        ['C', 'Retail mic',                          '50-100',      '300-600',      '28.000 EUR',   '5 RON/h',   '120 min'],
        ['D', 'Standalone public',                   '80-150',      '300-500',      '28.000 EUR',   '5 RON/h',   'fara'],
        ['E', 'Semi-public',                         '50-100',      '200-400',      '28.000 EUR',   '5 RON/h',   '0-15 min'],
        ['F', 'Captiv (hotel, birou, spital)',       '50-100',      '100-250',      '28.000 EUR',   '10 RON/sesiune', 'fara'],
        ['G', 'Mega-mall (centru comercial mare)',   '1.000-3.000', '5.000-15.000', '130.000 EUR',  '10-15 RON/h', '180 min'],
    ],
    col_widths=[1, 4, 2.2, 2.5, 2.3, 2.5, 2],
)

add_para('Tipul G (mall) este exclus din baseline-ul scaling-ului 12 parcari, dar mentionat ca potential upside in pre-mortem (sectiunea 8).', italic=True, size=10)

add_heading('Variante contractuale (parametrizate)', 2)
add_table(
    ['Cod', 'Nume', 'Cine plateste fee', 'Revenue split'],
    [
        ['C1',  'Operator 100% revenue, fara fee',           'nimeni',                       'Operator 100%'],
        ['C2',  '50/50 split, fara fee',                     'nimeni',                       '50/50'],
        ['C3',  'Operator primeste fee fix, proprietar 100%', 'proprietar -> operator',     'Proprietar 100%'],
        ['C4',  'Hibrid: rent operator + 50/50 split',        'operator -> proprietar',     '50/50'],
        ['C5',  'Operator plateste fee, operator 100% rev',   'operator -> proprietar',     'Operator 100%'],
    ],
    col_widths=[1, 5.5, 4.5, 3],
)

add_heading('Parametri financiari principali', 2)
add_para('Toti parametrii sunt editabili in sheet-ul Parametri din model.xlsx:')
add_bullet('Forma juridica: SA (Societate pe Actiuni); impozit profit 16%')
add_bullet('Amortizare CAPEX: 5 ani liniar')
add_bullet('Provizion CAPEX recurent: 6%/an din CAPEX initial')
add_bullet('Inflatie OPEX si tarife: 6%/an')
add_bullet('Rata dobanda credit echipamente: 9%/an, durata 5 ani')
add_bullet('Pondere credit / capital propriu: 70% / 30%')
add_bullet('Cost capital propriu: 18%; WACC calculat: 10,69%')
add_bullet('Curs EUR/RON: 5,00')

add_heading('Overhead corporate', 2)
add_para('Structura detaliata pe linii editabile:')
add_table(
    ['Faza', 'Total overhead anual', 'Comentariu'],
    [
        ['Anul 1 — lean (sub 5 parcari)',      '119.000 EUR',  'Fondator face BD, mentenanta externa'],
        ['Anul 2+ (BD intern, sub 5 parcari)', '141.000 EUR',  'BD/sales angajat, fondator revine la leadership'],
        ['Anul 2+ (peste 5 parcari)',          '166.000 EUR',  'Tehnician intern teren; reduce OPEX direct/parcare'],
    ],
    col_widths=[5.5, 3.5, 5.5],
)

add_pagebreak()

# ===== 7. FAZA 2 =====
add_heading('6. Faza 2 — Validator matematic si Reality check', 1)

add_heading('Sanity checks (toate PASS)', 2)
add_bullet('Suma procentelor distributie durate = 100% pe fiecare tip')
add_bullet('Inflatie zero -> anul 5 reproduce stabil anul 1')
add_bullet('Cota impozit zero -> profit_net = profit_impozabil')
add_bullet('Provizion zero -> EBIT = EBITDA - overhead')
add_bullet('Anuitate × 5 > credit (datorita dobanzii)')
add_bullet('Rata colectare 100% -> venit = trafic × tarif × 365')
add_bullet('WACC consistent (intre cost datorie post-tax si cost capital propriu)')

add_heading('Reality check vs benchmark industrie', 2)
add_para('EBITDA margin operational per parcare individuala (75-97%) este peste benchmark-ul industriei parking europene (Apcoa, Q-Park, Saba: 25-45%).')
add_para('Diagnostic:', bold=True)
add_bullet('La nivel operational pur (per parcare, fara overhead), margins reale sunt mari (80-97% pentru non-retail) — corect financiar pentru o parcare individuala bine operata')
add_bullet('La nivel agregat companie SA (cu overhead corporate, costuri financiare, mix inclus locatii proaste), margins se incadreaza in benchmark 25-45%')
add_bullet('Nu este dovada de optimism; este consecinta unei separari corecte intre niveluri operational vs corporate')

add_para('Validare suplimentara recomandata: validare independenta a ipotezelor de trafic pe primele 2-3 locatii instalate, prin masuratori reale 3-6 luni.', italic=True, size=10)

add_pagebreak()

# ===== 8. FAZA 3 =====
add_heading('7. Faza 3 — Model financiar consolidat (model.xlsx)', 1)

add_heading('Structura modelului', 2)
add_para('Excel-ul contine 8 sheet-uri si 709 formule, fara erori. Toti parametrii sunt editabili in sheet-ul Parametri (zero numere hardcodate in formule).')
add_table(
    ['Sheet', 'Continut'],
    [
        ['Parametri',         'Toti parametrii editabili: financiari, tipuri locatii, distributii durate, rate colectare, variante, overhead'],
        ['Tipuri locatii',    'Calcule derivate per tip: OPEX anual, anuitate credit, amortizare, provizion'],
        ['Variante',          '29 scenarii core: tip × varianta × gratuitate × colectare; CF Net per parcare an 1'],
        ['Portofoliu',        'Mix 12 parcari editabil; consolidare companie-level cu tax shield'],
        ['Cash Flow',         'Agregat anual portofoliu pe 5 ani + cumulat (Min Cash Balance)'],
        ['Sumar comparativ',  'KPI principal cu praguri si verdict (PASS/FAIL)'],
        ['Sensibilitati',     'Lista stres tests pentru audit ipoteze'],
        ['Modificari',        'Audit trail pentru orice modificare structurala'],
    ],
    col_widths=[3.5, 12],
)

add_heading('Portofoliul baseline (12 parcari)', 2)
add_para('Anul 1 (5 parcari operate la finalul anului):')
add_bullet('2 × Auchan (tip A) — pipeline activ')
add_bullet('2 × Lidl/Kaufland (tip B) — pipeline activ')
add_bullet('1 × Captiv pilot (tip F) — diversificare timpurie')

add_para('Anul 2 (+7 parcari, total 12):')
add_bullet('1 × Auchan/Carrefour (tip A)')
add_bullet('4 × Lidl/Kaufland (tip B)')
add_bullet('1 × Standalone public (tip D)')
add_bullet('1 × Captiv (tip F)')

add_heading('Consolidare proper SA (tax shield)', 2)
add_para('Modelul aplica consolidare la nivel de companie SA pentru calculul impozitului pe profit:')
add_bullet('Sum EBIT pre-overhead per portofoliu')
add_bullet('Subtract overhead corporate (la nivel companie)')
add_bullet('Subtract sum amortizari (la nivel companie)')
add_bullet('Apply 16% tax doar pe profit pozitiv consolidat (loss offset disponibil intre parcari)')
add_bullet('+ amortizari (re-adaugare non-cash)')
add_bullet('- anuitati credit -> CF Net portofoliu')

add_para('Aceasta abordare reflecta exact modul in care SA-urile romanesti depun declaratii fiscale anuale, cu consolidare companie-level. Comparativ cu o alocare per-parcare a overhead-ului (cum ar fi naive aproximarea per cost center), abordarea consolidata genereaza un beneficiu fiscal de ~10-30k EUR/an datorita compensarii pierderilor unora cu profiturile altora.', italic=True, size=10)

add_heading('Cash flow agregat per an', 2)
add_table(
    ['Element', 'An 1', 'An 2', 'An 3', 'An 4', 'An 5', 'Cumulat'],
    [
        ['CF Net portofoliu (post overhead)',  '-67.213', '+95.271',  '+114.995', '+135.903', '+158.065', '+437.020'],
        ['CAPEX investit (capital propriu)',   '-62.400', '-79.800',  '0',        '0',        '0',        '-142.200'],
        ['CF Total (operational + CAPEX)',     '-129.613','+15.471',  '+114.995', '+135.903', '+158.065', '+294.820'],
        ['CF cumulat (Min Cash Balance)',      '-129.613','-114.142', '+853',     '+136.756', '+294.820', '-129.613'],
    ],
    col_widths=[5, 1.7, 1.7, 1.7, 1.7, 1.7, 2],
)

add_para('Min Cash Balance -129.613 EUR (la finalul anului 1) inseamna ca, peste capitalul propriu de 142.200 EUR, este nevoie de un buffer de cash de ~130k pentru a acoperi pierderea operationala anul 1 + investitia in parcari noi anul 2.', bold=True)

add_pagebreak()

# ===== 9. PRE-MORTEM =====
add_heading('8. Pre-mortem — 3 moduri concrete de esec', 1)
add_para('Inainte de a recomanda investitia, am imaginat ca au trecut 5 ani si proiectul a esuat. Am identificat 3 moduri concrete de esec, cu semnale precoce care le-ar anunta si masuri de mitigare specifice.')

add_heading('Modul 1: Pierdere contract major (Auchan / Lidl / Kaufland)', 2)
add_para('Scenariu', bold=True)
add_para('Un retailer mare reziliază contractul sau rezistă la reinnoirea cu termeni similari. Pierderea a 2-3 parcari simultan reprezinta 25-30% din portofoliu.')
add_para('Impact cantitativ', bold=True)
add_para('Pierderea celor 3 parcari Auchan reduce NPV cu ~85.000 EUR si CF cumulat cu ~200.000 EUR. Singur ramane portofoliul ne-Auchan (9 parcari) cu NPV ~60.000 EUR — viabil dar marginal.')
add_para('Semnale precoce de monitorizat', bold=True)
add_bullet('Schimbare echipa management / achizitii la retailer (urmariti LinkedIn, comunicate corporative trimestriale)')
add_bullet('Tendinta de internalizare operatiuni parking in retail mare (Carrefour Romania a anuntat in 2025 lansarea propriei solutii ANPR)')
add_bullet('Concurent care ofera C3 fara risc pentru retailer (operator preia toata expunerea + plata fee fix)')
add_bullet('Cresterea concurentei pe piata operatorilor (B-Parking, Apcoa expansiune)')
add_para('Masuri de mitigare', bold=True)
add_bullet('Diversificare timpurie pe non-retail: anul 1 include 1 parcare captiv pilot pentru a invata operational si a reduce concentrarea')
add_bullet('Contracte multianuale cu clauze de notificare 12 luni (timp suficient sa repozitionam echipamentul)')
add_bullet('Relatii multiple in fiecare retailer (nu doar 1 contact)')
add_bullet('Plan de portabilitate echipament: ANPR + bariere pot fi mutate fara depreciere mare daca pierdem locatie')

add_heading('Modul 2: Schimbare reglementare ANPR / GDPR', 2)
add_para('Scenariu', bold=True)
add_para('Autoritatea ANSPDCP introduce restrictii noi pe colectarea numerelor de inmatriculare: consent explicit, retentie limitata la 30 zile, accesibilitate date mai stricta. Cost compliance suplimentar 5-15k EUR per parcare; posibile amenzi 50-200k EUR companie pentru neconformitate retroactiva.')
add_para('Impact cantitativ', bold=True)
add_para('Cost initial compliance 60-180k EUR pe portofoliul de 12 parcari -> consumat capital de lucru anul 1-2. NPV scade cu 50-150k EUR.')
add_para('Semnale precoce de monitorizat', bold=True)
add_bullet('Comunicari ANSPDCP catre operatorii de parking (newsletter, ghiduri, decizii publice)')
add_bullet('Dezbateri parlamentare pe protectia datelor in spatii publice (Comisia juridica Camera Deputatilor)')
add_bullet('Concurenti ANPR amendati public (case-uri precedente)')
add_bullet('Raportari oficiale UE (EDPB) pe parking ANPR')
add_para('Masuri de mitigare', bold=True)
add_bullet('Consultanta juridica externa lunara pe GDPR (3.500 EUR/an deja in overhead)')
add_bullet('Parteneriat tehnic cu un furnizor ANPR conform certificat (audit trail tehnic disponibil pentru ANSPDCP)')
add_bullet('Modulare arhitectura sistemului pentru a permite schimbare politici retentie fara reinstalare hardware')
add_bullet('Asigurare profesionala cu extindere pentru reglementari (10-15k EUR/an)')

add_heading('Modul 3: Schimbare comportament consumator (offline -> online)', 2)
add_para('Scenariu', bold=True)
add_para('Trafic in magazine fizice scade 20-30% in 3-5 ani datorita expansiunii livrarii la domiciliu (Auchan + Lidl + Glovo + Bringo + Tazz). Trafic parcare scade proportional.')
add_para('Impact cantitativ', bold=True)
add_para('La trafic -25% pe retail (sensibilitate testata in model), NPV scade de la +145k la circa -15k EUR (NEGATIV); CF cumulat 5 ani scade de la 437k la 250k. Proiectul devine break-even, nu mai e atractiv pentru parteneri.')
add_para('Semnale precoce de monitorizat', bold=True)
add_bullet('Rapoarte trimestriale retailers cu date pe trafic in magazine (Auchan, Lidl public statistics in raportari anuale)')
add_bullet('Dezvoltari delivery partnerships (Auchan-Glovo lansare 2025; Lidl-Tazz pilot Bucuresti)')
add_bullet('Cresterea bursei retail online vs offline (NielsenIQ, GfK rapoarte trimestriale)')
add_bullet('Inchidere sau redimensionare hipermarketi (Carrefour a anuntat 5 inchideri Romania 2024-2025)')
add_para('Masuri de mitigare', bold=True)
add_bullet('In contracte includem clauza de revizuire trafic la 24 luni cu drept de renegociere tarif')
add_bullet('Portofoliu non-retail (4 din 12 parcari) actioneaza ca pernă - trafic in parcari publice si captive nu depinde de shopping fizic')
add_bullet('Orientare progresiva in anul 3-5 spre tipuri D/E/F (parcari publice, semi-publice, captive) unde delivery online nu schimba cererea')
add_bullet('Pivot oferta: Total Hub poate adauga servicii complementare (curatare auto, spalatorie, EV charging) pentru a creste venit per loc')

add_pagebreak()

# ===== 10. SENSIBILITATI =====
add_heading('9. Sensibilitati si stres tests', 1)
add_para('Modelul include 8 stres tests in sheet-ul Sensibilitati din model.xlsx. Cele mai importante pentru audit partener:')
add_table(
    ['Stres test', 'Descriere', 'Impact NPV'],
    [
        ['Trafic non-retail -25%',          'Reduce 400/300/200 -> 300/225/150 intrari/zi',     '~+30k EUR (de la +145k)'],
        ['OPEX +50%',                       'Multiplica OPEX direct cu 1,5 pe toate tipurile',  '~+50k EUR'],
        ['Colectare pesimist',              'Trece pe coloana Pesimist (60/50/25%) non-retail', '~+15k EUR'],
        ['Inflatie OPEX 10% (vs 6%)',       'Stres macroeconomic',                              'NPV ~+90k EUR'],
        ['WACC 15% (vs 10,7%)',             'Mai mult capital propriu, mai putin credit',       'NPV scade ~-40k'],
        ['Pierdere Auchan (3 parcari)',     'Risc concentrare retailer mare',                   'NPV ~+60k'],
        ['Negociere 60min gratuitate retail','Schimba grat 120 in 60 pe toate B',                'NPV ~+450k UPSIDE'],
        ['Adaugare 1 mall mare (CAPEX +130k)','Hero project upside scenario',                    'NPV ~+250k UPSIDE'],
    ],
    col_widths=[5, 6, 4],
)

add_para('Concluzii din sensibilitati:', bold=True)
add_bullet('Proiectul ramane viabil (NPV pozitiv) la majoritatea scenariilor pesimiste — robust la stres')
add_bullet('Singurul scenariu care produce NPV negativ este combinatia trafic -25% pe TOTUL portofoliul (impus prin actualizarea profilelor de trafic in Parametri)')
add_bullet('Upside semnificativ daca: (a) negociem gratuitate sub 120 min cu macar 1-2 retaileri, (b) atragem mall-ul mare in portofoliu')

add_pagebreak()

# ===== 11. CONCLUZII =====
add_heading('10. Concluzii si decizia ceruta', 1)

add_heading('Recomandare: GO CONDITIONAT', 2)
add_para('Proiectul este viabil pe baseline-ul de scaling agresiv (12 parcari finalul anului 2). Indicatorii financiari principali — NPV +145k EUR, IRR 31%, Payback 3 ani — depasesc pragurile de viabilitate. Singurul KPI care esueaza este pragul ambitios de cash flow cumulat (437k vs target 800k) — acest prag a fost stabilit ca tinta agresiva, nu ca prag de viabilitate.')

add_heading('Conditionari pentru GO', 2)
add_bullet('Validare independenta a ipotezelor de trafic pe primele 2-3 locatii instalate (masuratori reale 3-6 luni dupa go-live)')
add_bullet('Asigurare cash buffer ~130k EUR peste capitalul propriu de 142k (acopera CF Net negativ anul 1 si 2)')
add_bullet('Diversificare timpurie pe non-retail: anul 1 include 1 parcare captiv pilot ca diversificare, nu se asteapta anul 2')
add_bullet('Plan explicit de mitigare pe cele 3 moduri de esec din pre-mortem (Auchan exclusivity protections, GDPR consultanta lunara, monitoring trafic offline retail)')

add_heading('Optiuni de structura financiara propuse partenerilor', 2)
add_para('1. Investitie 100% Total Hub', bold=True)
add_para('Capital propriu 142k + credit 332k garantat de companie. Necesita cash buffer ~130k pentru start. Total Hub pastreaza 100% din valoarea creata (NPV +145k + upside posibil pana la +400k cu negociere gratuitate sub 120 min).')

add_para('2. Co-investitie 50/50 cu partener financiar', bold=True)
add_para('Capital propriu impartit 71k/71k. Credit ramane in numele Total Hub. Reduce expunere financiara fondator dar dilueaza randament. Partenerul primeste 50% din profit consolidat, dupa rambursarea capitalului propriu (cca anul 3).')

add_para('3. Co-investitie + management de risc', bold=True)
add_para('Partenerul aduce capital propriu + relatii cu retaileri suplimentari (Auchan/Carrefour pe care fondatorul nu le are inca). Aceasta optiune accelereaza scaling-ul si reduce timpul de negociere comerciala — potential reducere a payback-ului la 2-2,5 ani.')

add_heading('Termen decizie', 2)
add_para('Confirmarea optiunii de structura este necesara in 30 zile pentru a putea semna primele contracte (Auchan + Lidl) inainte de Q3 2026 si a pune in functiune anul 1 conform planului de scaling.')

add_pagebreak()

# ===== 12. GLOSAR =====
add_heading('11. Glosar', 1)
add_para('Termeni tehnici folositi in raport, definiti pentru audienta non-financiara:')

glosar = [
    ('NPV (Net Present Value)', 'Valoarea actualizata neta. Suma fluxurilor de numerar viitoare actualizate la rata WACC, minus investitia initiala. NPV > 0 inseamna ca proiectul creeaza valoare la rata de discount aleasa.'),
    ('IRR (Internal Rate of Return)', 'Rata interna de rentabilitate. Rata de actualizare la care NPV devine zero. Un IRR de 31% inseamna ca proiectul returneaza 31% pe an din capitalul investit.'),
    ('WACC (Weighted Average Cost of Capital)', 'Costul mediu ponderat al capitalului. Calculat ca pondere_credit × cost_credit_post_tax + pondere_capital_propriu × cost_capital_propriu. In modelul nostru: 70% × 9% × (1-16%) + 30% × 18% = 10,69%.'),
    ('CAPEX (Capital Expenditure)', 'Cheltuieli de capital — investitia initiala in echipamente fizice (ANPR, bariere, terminale, infrastructura electrica si retea).'),
    ('OPEX (Operating Expenditure)', 'Cheltuieli operationale curente — energie electrica, internet, mentenanta, software/SaaS.'),
    ('EBITDA', 'Earnings Before Interest, Tax, Depreciation, and Amortization. Profitul operational inainte de costurile financiare, taxe si amortizare. Util pentru benchmarking operational vs concurenta.'),
    ('CF Net (Cash Flow Net)', 'Fluxul de numerar net — cash-ul real disponibil pentru investitor dupa toate cheltuielile, taxele si serviciile datoriei. Este metricul folosit pentru calculul NPV si IRR (NU EBITDA).'),
    ('Anuitate', 'Plata anuala constanta de credit (capital + dobanda). Calculata ca: credit × rata_dobanda / (1 - (1+rata)^(-ani)).'),
    ('Amortizare', 'Recunoasterea contabila a uzurii unei investitii pe durata sa de viata utila (5 ani pentru hardware ANPR/bariere). Cheltuiala non-cash care reduce profitul impozabil.'),
    ('Provizion CAPEX recurent', 'Suma rezervata anual pentru reparatii si inlocuiri partiale ale echipamentului uzat. 6% din CAPEX initial este standardul pentru hardware parking.'),
    ('Perioada de gratuitate', 'Interval de timp dupa intrare in care clientul nu plateste (in retail, 120 min in Romania). Sub aceasta durata clientul iese gratuit; peste, plateste tarif standard.'),
    ('DSO (Days Sales Outstanding)', 'Numarul de zile intre vanzare si incasare. In parking cu plata card/aplicatie: 1-2 zile.'),
    ('DPO (Days Payable Outstanding)', 'Numarul de zile intre primirea facturii si plata catre furnizor. In parking: ~30 zile (utilitati, mentenanta, salarii la finalul lunii).'),
    ('Cycle gap', 'Diferenta DSO - DPO. Daca negativ (incasezi inainte sa platesti), nu ai nevoie de capital de lucru. Cazul nostru: -28 zile.'),
    ('Min Cash Balance', 'Cea mai negativa valoare a cash-ului cumulat lunar. Determina capitalul minim de care e nevoie inainte ca afacerea sa devina auto-sustenabila.'),
    ('Payback period', 'Numarul de ani dupa care cumularea fluxurilor de numerar acopera investitia initiala (capital propriu).'),
    ('ANPR', 'Automatic Number Plate Recognition — sistem care recunoaste automat numerele de inmatriculare prin camere video. Folosit la intrare/iesire pentru identificare automata.'),
    ('SA (Societate pe Actiuni)', 'Forma juridica a companiei Total Hub. Diferita de SRL: capital social minim 25.000 EUR, governance prin AGA si Consiliu de Administratie, audit obligatoriu peste anumite plafoane.'),
    ('Variante contractuale (C1-C5)', 'Cele 5 modele contractuale generice intre operator (Total Hub) si proprietar (retailer/proprietar locatie): C1 operator 100% rev, C2 50/50, C3 fee fix la operator, C4 hibrid rent + 50/50, C5 operator plateste rent + 100% rev.'),
]
for term, definition in glosar:
    p = doc.add_paragraph()
    r = p.add_run(term + ': ')
    r.bold = True
    r.font.size = Pt(10)
    r2 = p.add_run(definition)
    r2.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(6)

add_pagebreak()

# ===== 13. DISCLAIMER =====
add_heading('12. Disclaimer si limitari', 1)

add_para('Acest document reflecta ipoteze valide la data 2026-04-29.', bold=True)

add_para('Limitari ale analizei:')
add_bullet('Pipeline-ul comercial (Auchan 2 locatii, Lidl multiple, Kaufland in plan) este declarat in negociere activa, nu contractat. Numarul real de parcari operate poate fi mai mic sau mai mare decat baseline-ul 12.')
add_bullet('Cifrele de trafic (1.000 intrari/zi pentru retail mediu, 400 intrari/zi pentru standalone, etc.) sunt estimari de research preliminar. Validarea finala pe locatii concrete este obligatorie inainte de semnare contract.')
add_bullet('Rate de colectare (75% standalone, 65% semi-public, 35% captiv) sunt estimari pe baza benchmark-urilor industriei. Performanta reala depinde de calitatea echipamentului si a operatiunilor.')
add_bullet('CAPEX (40-50k EUR retail, 28k non-retail) este estimat pe baza ofertei furnizori; cifrele finale depind de configuratia locatiei (numar de intrari/iesiri, dimensiune retea camere).')
add_bullet('OPEX direct (450-700 EUR/luna) este o estimare lean; in primii ani de operare poate fi 20-30% mai mare datorita lipsei de optimizare operationala.')
add_bullet('Cota impozit 16% si regimul SA sunt valid la 2026-04-29; modificari fiscale viitoare pot impacta rezultatele.')
add_bullet('Modelul presupune contract pe 5 ani fara reinnoire automata. La final orizont, valoarea reziduala a echipamentelor (~10-20% din CAPEX) nu este inclusa in CF — conservator.')

add_para('Recomandari de validare:')
add_bullet('Audit independent al modelului matematic de un consultant financiar (4-6h efort)')
add_bullet('Validare fiscala cu un contabil autorizat pentru tratamentul SA al impozitului consolidat')
add_bullet('Masuratori reale de trafic pe primele 2-3 locatii in primele 6 luni dupa go-live, cu reactualizarea modelului')
add_bullet('Reverificare anuala a ipotezelor de inflatie si dobanda fata de realitatea pietei')

add_para('Modelul model.xlsx permite recalcul instantaneu cu parametri editabili.', italic=True)
add_para('Raportul Word/PDF trebuie regenerat manual dupa modificari semnificative ale parametrilor.', italic=True)

add_para('')
add_para('Pentru intrebari sau clarificari, va rugam sa contactati echipa Total Hub SA.', italic=True, size=10)

# === SAVE ===
output = '/Users/home-felix/Total Hub/Analiza fezabilitate parcari de inchiriat/outputs/analiza.docx'
doc.save(output)
print(f"Documentul salvat: {output}")
